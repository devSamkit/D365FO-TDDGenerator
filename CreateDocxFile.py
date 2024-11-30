# Create document

from docx import Document
from TDDConstant import NA, TDDSTYLE
import os

def initDocument(templatePath: str):
    initDocument.document = Document(templatePath)

def addNone():
    initDocument.document.add_paragraph(NA)

def addLineBreak():
    initDocument.document.add_paragraph('')

def addParagraph(stringVal: str):
    initDocument.document.add_paragraph(stringVal)
    
def addParagraph_italic(constantValue: str, stringVal: str):
    para = initDocument.document.add_paragraph()
    para.add_run(constantValue).italic = True
    para.add_run().add_break()
    para.add_run(stringVal)

def addHeading(text: str, level: int):
    initDocument.document.add_heading(text, level)
    
def addHeading_num(stringVal, headerStyle):
    prag = initDocument.document.add_paragraph(stringVal)
    prag.style = initDocument.document.styles[headerStyle]

def addTable(tableDist):
    if(bool(tableDist)):
        counter = 0
        headerList = tableDist.pop(0)
        columns = len(headerList)
        table = initDocument.document.add_table(1, columns)
        table.style = TDDSTYLE
        hdr_cells = table.rows[0].cells
        for i in headerList:
            hdr_cells[counter].text = str(i)
            counter += 1

        rowsList = [[i for i in tableDist[x]] for x in tableDist.keys()]
        for rowsValue in rowsList:
            row_cells = table.add_row().cells
            rowCounter = 0
            for cellValue in rowsValue:
                if(cellValue is None):
                    cellValue = ''
                row_cells[rowCounter].text = cellValue
                rowCounter += 1
    else:
        addNone()

def addPageBreak():
    initDocument.document.add_page_break()

def saveDocx(docxPath, fileName):
    if(os.path.exists(docxPath) == False):
        os.makedirs(docxPath)
        
    filePath = os.path.join(docxPath, fileName)
    initDocument.document.save(filePath)
    # Open file
    os.startfile(filePath)